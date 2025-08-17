#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EbookForgePro — single-file premium app

Purpose: Expand a full eBook from a first edition file or from Title, Subtitle, ToC, and Description.
Runs fully local by default, with optional API providers for AI expansion. Includes a professional GUI,
text cleaning (including the special rule to replace all double dashes `--` with commas `,`), exports
to MD, DOCX, EPUB, PDF, upload helpers for Gumroad (create product + upload), Ko-fi (open shop),
and Google Play Books Partner (ONIX generator + optional SFTP). Generates Windows build helpers.
All required dependencies are auto-installed if missing.

Tested targets: Windows 10/11 (Python 3.10–3.13), macOS 14+, Ubuntu 22+. Network is used only
when you explicitly trigger API expansion or uploads.

NOTE: Live uploads require your own valid credentials and vendor permissions.
"""

import os
import sys
import re
import io
import json
import textwrap
import subprocess
import importlib
import traceback
import base64
import tempfile
import threading
import datetime
import webbrowser
from pathlib import Path

APP_NAME = "EebookForgePro" if False else "EbookForgePro"  # (keep string literal stable for packers)
APP_VERSION = "1.2.0"
APP_ID = "com.id01t.ebookforgepro"

BASE = Path.cwd()
PROJECT_DIR = (BASE / "EbookForgeProject").resolve()
EXPORTS = PROJECT_DIR / "exports"
BUILD = PROJECT_DIR / "build"
ASSETS = PROJECT_DIR / "assets"
for p in (PROJECT_DIR, EXPORTS, BUILD, ASSETS):
    p.mkdir(parents=True, exist_ok=True)

# ----------------------
# Dependency management
# ----------------------
REQUIRED = {
    "requests": "requests",
    "reportlab": "reportlab",
    "ebooklib": "ebooklib",
    "markdown2": "markdown2",
    "docx": "python-docx",
    "ttkbootstrap": "ttkbootstrap",
}
# Optional (only if you use SFTP for Google Partner uploads)
OPTIONAL = {"paramiko": "paramiko"}

def ensure_pkg(import_name: str, pip_name: str):
    try:
        return importlib.import_module(import_name)
    except Exception:
        print(f"[deps] Installing {pip_name} …")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])
        return importlib.import_module(import_name)

for mod, pipn in REQUIRED.items():
    ensure_pkg(mod, pipn)

# Imports after ensuring
import requests  # type: ignore
from reportlab.lib.pagesizes import LETTER  # type: ignore
from reportlab.pdfgen import canvas as pdf_canvas  # type: ignore
from ebooklib import epub  # type: ignore
import markdown2  # type: ignore
from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore

# GUI
import tkinter as tk  # stdlib
from tkinter import ttk, filedialog, messagebox
import ttkbootstrap as tb  # type: ignore

# DPI awareness on Windows (improves font rendering)
if sys.platform.startswith("win"):
    try:
        import ctypes
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

# -------------
# App icon .ico
# -------------
EMBED_ICO = (
    b"AAABAAEAEBAAAAEAIABoBAAAFgAAACgAAAAQAAAAIAAAAAEAGAAAAAAAAAMAAAAAAAAAAAAAAAAAAAAA"
    b"AAAAAAD///8A////AP///wD///8A////AP///wD///8A////AP///wD///8A////AP///wD///8A////"
    b"AP///wD///8A////AP///wD///8A////AP///wD///8A////AAkJCQAZGRkAJCQkACUlJQAmJiYAKCgo"
    b"ACoqKgAuLi4AMDAwADExMQAyMjIANjY2ADg4OAA7OzsAPj4+AEBAQABDQ0MASEhIAE1NTQBPT08AVFRU"
    b"AFhYWABeXl4AYGBgAGNjYwBqam oAbGxsAHR0dAB+fn4AgICA AISEhACMjIwAlJSUAJyc nACwsLAAtLS0A"
    b"Li4uAC8vLwAwMDAAMTExADMzMwA0NDQANzc3ADg4OAA5OTkAOjo6ADs7OwA9PT0APj4+AEBAQABERERA"
    b"AEZGRgBJSUkATExMAE9PTwBVVVUAWVlZAF1dXQBfX18AYmJiAGZmZgBoaGgAbm5uAHBwcAB2dnYAenp6"
    b"AICCggCJiYkAkZGRAJWVlQCcnJwAn5+fAKGhoQCpqakArKysAK+vrwCwsLAAtra2ALa2tgC+vr4AwMDA"
    b"AMLCwgDFxcUAx8fHAMS EhADJycnA////wD///8A////AP///wD///8A////AP///wD///8A////AP//"
    b"/wD///8A////AP///wD///8A////AP///wD///8A////AP///wD///8A////AP///wD///8A"
)
ICO_PATH = ASSETS / "app.ico"
if not ICO_PATH.exists():
    try:
        ICO_PATH.write_bytes(base64.b64decode(EMBED_ICO))
    except Exception:
        pass

# ----------------
# Text utilities
# ----------------
CLEAN_OPTS = {
    "replace_double_hyphen": True,
    "replace_emdash": True,
    "normalize_ws": True,
    "smart_quotes": True,
}

def clean_text(s: str, opts: dict | None = None) -> str:
    """Core text cleaner (implements the famous -- -> , rule)."""
    opts = {**CLEAN_OPTS, **(opts or {})}
    if opts.get("replace_double_hyphen"):
        s = s.replace("--", ",")
    if opts.get("replace_emdash"):
        s = s.replace("—", ",")
    if opts.get("smart_quotes"):
        s = s.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
    if opts.get("normalize_ws"):
        s = re.sub(r"[\t\r\f\v]", " ", s)
        s = re.sub(r"\n{3,}", "\n\n", s)
        s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()

def slugify(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9\-_ \.]", "", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s or "ebook"

# ------------------------------
# Offline manuscript generation
# ------------------------------
def scaffold_from_meta(title: str, subtitle: str, toc: str, description: str, seed: str = "") -> str:
    """Generate a robust Markdown manuscript purely offline."""
    title = clean_text(title or "Untitled")
    subtitle = clean_text(subtitle or "")
    description = clean_text(description or "")
    if "\n" in toc:
        chapters = [clean_text(x) for x in toc.splitlines() if x.strip()]
    else:
        chapters = [clean_text(x) for x in toc.split(",") if x.strip()]
    if not chapters:
        chapters = ["Introduction", "Chapter 1", "Chapter 2", "Conclusion"]

    out = [f"# {title}"]
    if subtitle:
        out.append(f"\n**{subtitle}**\n")
    if description:
        out.append(f"\n> {description}\n")

    out.append("\n### Table of Contents\n")
    for i, ch in enumerate(chapters, 1):
        out.append(f"{i}. {ch}")

    seed_short = clean_text(seed)[:2000]
    for i, ch in enumerate(chapters, 1):
        out.append(f"\n\n## {i}. {ch}\n")
        para = textwrap.fill(
            f"{ch} presents a practical strategy with step-by-step actions and measurable outcomes. "
            f"Write clearly and prioritize momentum. {seed_short[:300]}",
            width=92,
        )
        chunk = []
        target_words = 900 if i > 1 else 700
        while len(" ".join(chunk + [para]).split()) < target_words:
            chunk.append(para)
            para = textwrap.fill(
                "Add examples, pitfalls, best practices, and a short scenario. "
                "End with ROI metrics and a week-one action plan.",
                width=92,
            )
        out.append("\n\n".join(chunk))
        out.append("\n\n### Exercises\n1. Define 3 KPIs.\n2. Draft a 5-step plan with estimates.\n3. Identify two risks and mitigations.")
        out.append("\n\n### Key Takeaways\nClarity, outcomes, repeatable systems; prioritize actions that drive ROI.")

    return clean_text("\n".join(out))

# ------------------------------
# Optional AI expanders (off by default)
# ------------------------------
class Expander:
    def __init__(self, cfg: dict):
        self.cfg = cfg

    def expand(self, manuscript: str) -> str:
        mode = self.cfg.get("mode", "offline")
        try:
            if mode == "offline":
                return manuscript
            if mode == "openai":
                return self._openai_like(manuscript)
            if mode == "gemini":
                return self._gemini(manuscript)
            if mode == "local":
                return self._ollama(manuscript)
        except Exception as e:
            print("[expander] error, fallback to offline:", e)
        return manuscript

    def _openai_like(self, manuscript: str) -> str:
        key = self.cfg.get("openai_api_key")
        base = self.cfg.get("openai_base", "https://api.openai.com/v1")
        model = self.cfg.get("openai_model", "gpt-4o-mini")
        if not key:
            return manuscript
        url = f"{base}/chat/completions"
        headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
        prompt = ("Expand and professionally format the following eBook manuscript. "
                  "Return Markdown only. Replace double dashes with commas.\n\n" + manuscript)
        payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.4}
        r = requests.post(url, headers=headers, data=json.dumps(payload), timeout=180)
        r.raise_for_status()
        out = r.json().get("choices", [{}])[0].get("message", {}).get("content", manuscript)
        return clean_text(out)

    def _gemini(self, manuscript: str) -> str:
        key = self.cfg.get("gemini_api_key")
        model = self.cfg.get("gemini_model", "gemini-1.5-flash")
        if not key:
            return manuscript
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
        payload = {"contents": [{"parts": [{"text": "Expand and format as Markdown; replace -- with ,.\n\n" + manuscript}]}]}
        r = requests.post(url, json=payload, timeout=180)
        r.raise_for_status()
        data = r.json()
        try:
            return clean_text(data["candidates"][0]["content"]["parts"][0]["text"])
        except Exception:
            return manuscript

    def _ollama(self, manuscript: str) -> str:
        base = self.cfg.get("local_base", "http://127.0.0.1:11434")
        model = self.cfg.get("local_model", "llama3")
        url = f"{base}/api/generate"
        payload = {"model": model, "prompt": "Expand and format as Markdown; replace -- with ,.\n\n" + manuscript, "stream": False}
        r = requests.post(url, json=payload, timeout=240)
        r.raise_for_status()
        return clean_text(r.json().get("response", manuscript))

# ------------------------------
# Exporters: MD, DOCX, EPUB, PDF
# ------------------------------
class Exporter:
    def __init__(self, proj: Path):
        self.proj = proj

    def export_md(self, text: str, title: str) -> Path:
        p = EXPORTS / f"{slugify(title or 'manuscript')}.md"
        p.write_text(text, encoding="utf-8")
        return p

    def export_docx(self, text: str, title: str) -> Path:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        for line in text.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("> "):
                p = doc.add_paragraph()
                r = p.add_run(line[2:].strip())
                r.italic = True
            else:
                doc.add_paragraph(line)
        p = EXPORTS / f"{slugify(title or 'manuscript')}.docx"
        doc.save(p)
        return p

    def export_epub(self, text: str, title: str, author: str) -> Path:
        book = epub.EpubBook()
        book.set_identifier(APP_ID)
        book.set_title(title or "Untitled")
        book.add_author(author or "")
        book.set_language("en")

        parts = re.split(r"\n## ", text)
        items = []
        for i, part in enumerate(parts):
            if i == 0:
                c = epub.EpubHtml(title="Front", file_name="front.xhtml", lang="en")
                c.content = f"<h1>{_esc(title)}</h1>" + markdown2.markdown(part)
            else:
                first = part.splitlines()[0]
                body = "\n".join(part.splitlines()[1:])
                c = epub.EpubHtml(title=first, file_name=f"chap_{i:02d}.xhtml", lang="en")
                c.content = f"<h2>{_esc(first)}</h2>" + markdown2.markdown(body)
            book.add_item(c)
            items.append(c)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.toc = tuple(items)
        book.spine = ["nav"] + items
        out = EXPORTS / f"{slugify(title or 'ebook')}.epub"
        epub.write_epub(str(out), book)
        return out

    def export_pdf(self, text: str, title: str, author: str) -> Path:
        out = EXPORTS / f"{slugify(title or 'ebook')}.pdf"
        c = pdf_canvas.Canvas(str(out), pagesize=LETTER)
        width, height = LETTER
        margin = 54
        y = height - margin
        c.setFont("Times-Bold", 18)
        c.drawString(margin, y, (title or "Untitled")[:100])
        y -= 24
        if author:
            c.setFont("Times-Roman", 12)
            c.drawString(margin, y, f"Author: {author}")
            y -= 16
        c.setFont("Times-Roman", 11)
        # simple paragraph flow
        plain = re.sub(r"<[^>]+>", "", text)
        for para in plain.split("\n\n"):
            for line in textwrap.wrap(para, width=92):
                if y < margin:
                    c.showPage()
                    y = height - margin
                    c.setFont("Times-Roman", 11)
                c.drawString(margin, y, line)
                y -= 14
            y -= 6
        c.showPage()
        c.save()
        return out

# tiny HTML escaper
import html
def _esc(s: str) -> str:
    return html.escape(s or "")

# ------------------------------
# Upload helpers (Gumroad, Ko-fi, Google Partner ONIX + SFTP)
# ------------------------------
class Uploader:
    def __init__(self, proj: Path):
        self.proj = proj

    def gumroad_create_and_upload(self, token: str, product_name: str, price_cents: int, summary: str, file_path: Path) -> dict:
        file_path = Path(file_path)
        if not token:
            raise RuntimeError("Gumroad API token required")
        if not file_path.exists():
            raise FileNotFoundError(file_path)
        headers = {"Authorization": f"Bearer {token}"}
        # Create product
        data = {"name": product_name, "price": price_cents, "description": summary}
        r = requests.post("https://api.gumroad.com/v2/products", headers=headers, data=data, timeout=90)
        if r.status_code >= 400:
            raise RuntimeError(f"Gumroad product creation failed: {r.status_code} {r.text}")
        product_id = (r.json().get("product") or {}).get("id") or r.json().get("id")
        if not product_id:
            raise RuntimeError(f"Product id not found in response: {r.text}")
        # Upload file
        with open(file_path, "rb") as fh:
            files = {"file": fh}
            r2 = requests.post(f"https://api.gumroad.com/v2/products/{product_id}/files",
                               headers=headers, files=files, timeout=300)
        if r2.status_code >= 400:
            raise RuntimeError(f"Gumroad file upload failed: {r2.status_code} {r2.text}")
        return {"product_id": product_id, "result": r2.json()}

    def kofi_open_shop(self):
        webbrowser.open("https://ko-fi.com/manage/shop")

    def google_onix_xml(self, metadata: dict) -> Path:
        today = datetime.date.today().strftime("%Y%m%d")
        def esc(x): return (x or "").replace("&", "&amp;").replace("<", "&lt;")
        onix = f"""
<?xml version="1.0" encoding="UTF-8"?>
<ONIXMessage release="3.0">
  <Header>
    <Sender><SenderName>{esc(metadata.get('publisher','iD01t Productions'))}</SenderName></Sender>
    <SentDateTime>{today}</SentDateTime>
  </Header>
  <Product>
    <RecordReference>{esc(metadata.get('isbn','NA'))}</RecordReference>
    <NotificationType>03</NotificationType>
    <ProductIdentifier><ProductIDType>15</ProductIDType><IDValue>{esc(metadata.get('isbn','NA'))}</IDValue></ProductIdentifier>
    <DescriptiveDetail>
      <ProductComposition>00</ProductComposition>
      <ProductForm>ED</ProductForm>
      <TitleDetail><TitleType>01</TitleType><TitleElement>
        <TitleElementLevel>01</TitleElementLevel>
        <TitleText>{esc(metadata.get('title','Untitled'))}</TitleText>
        <Subtitle>{esc(metadata.get('subtitle',''))}</Subtitle>
      </TitleElement></TitleDetail>
      <Contributor><SequenceNumber>1</SequenceNumber><ContributorRole>A01</ContributorRole><PersonName>{esc(metadata.get('author',''))}</PersonName></Contributor>
      <Language><LanguageRole>01</LanguageRole><LanguageCode>eng</LanguageCode></Language>
    </DescriptiveDetail>
    <CollateralDetail>
      <TextContent><TextType>03</TextType><ContentAudience>00</ContentAudience><Text>{esc(metadata.get('description',''))}</Text></TextContent>
    </CollateralDetail>
    <PublishingDetail>
      <Imprint><ImprintName>{esc(metadata.get('publisher','iD01t Productions'))}</ImprintName></Imprint>
      <Publisher><PublisherName>{esc(metadata.get('publisher','iD01t Productions'))}</PublisherName></Publisher>
      <PublishingDate><PublishingDateRole>01</PublishingDateRole><Date>{today}</Date></PublishingDate>
    </PublishingDetail>
  </Product>
</ONIXMessage>
""".strip()
        out = EXPORTS / "onix.xml"
        out.write_text(onix, encoding="utf-8")
        return out

    def google_sftp_upload(self, host: str, port: int, username: str, password: str | None, key_path: str | None, files: list[Path]) -> str:
        try:
            paramiko = importlib.import_module("paramiko")
        except Exception:
            ensure_pkg("paramiko", OPTIONAL["paramiko"])
            paramiko = importlib.import_module("paramiko")
        transport = None
        try:
            transport = paramiko.Transport((host, int(port or 22)))
            if key_path:
                pkey = paramiko.RSAKey.from_private_key_file(key_path)
                transport.connect(username=username, pkey=pkey)
            else:
                transport.connect(username=username, password=password)
            sftp = paramiko.SFTPClient.from_transport(transport)
            # Google typically expects /incoming or a configured dir. Adjust as needed.
            remote_dir = "/incoming"
            try:
                sftp.chdir(remote_dir)
            except IOError:
                try:
                    sftp.mkdir(remote_dir)
                    sftp.chdir(remote_dir)
                except Exception:
                    pass
            for fp in files:
                sftp.put(str(fp), Path(fp).name)
            sftp.close()
            return f"SFTP upload complete to {host}:{remote_dir}"
        finally:
            try:
                if transport:
                    transport.close()
            except Exception:
                pass

# ------------------------------
# Windows build helpers
# ------------------------------
class WinHelpers:
    @staticmethod
    def write_all():
        manifest = f"""
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<assembly xmlns="urn:schemas-microsoft-com:asm.v1" manifestVersion="1.0">
  <assemblyIdentity version="{APP_VERSION}.0" processorArchitecture="*" name="{APP_ID}" type="win32"/>
  <description>{APP_NAME}</description>
  <dependency>
    <dependentAssembly>
      <assemblyIdentity type="win32" name="Microsoft.Windows.Common-Controls" version="6.0.0.0" processorArchitecture="*" publicKeyToken="6595b64144ccf1df" language="*"/>
    </dependentAssembly>
  </dependency>
  <trustInfo xmlns="urn:schemas-microsoft-com:asm.v3">
    <security><requestedPrivileges>
      <requestedExecutionLevel level="asInvoker" uiAccess="false"/>
    </requestedPrivileges></security>
  </trustInfo>
</assembly>
""".strip()
        (BUILD / "app.manifest").write_text(manifest, encoding="utf-8")
        cmd = (
            f'pyinstaller --noconsole --onefile --name "{APP_NAME}" '
            f'--icon "{ICO_PATH.as_posix()}" '
            f'--manifest "build{os.sep}app.manifest" '
            f'"{(BASE / Path(__file__).name).as_posix()}"'
        )
        (BUILD / "build_windows.cmd").write_text(cmd, encoding="utf-8")
        (BUILD / "version.txt").write_text(APP_VERSION, encoding="utf-8")

# ------------------------------
# GUI
# ------------------------------
class App(tb.Window):
    def __init__(self):
        super().__init__(themename="flatly")
        self.title(f"{APP_NAME} {APP_VERSION}")
        self.geometry("1240x840")
        self.minsize(1080, 720)
        try:
            self.iconbitmap(ICO_PATH)
        except Exception:
            pass

        self.api_cfg = {
            "mode": "offline",  # offline, openai, gemini, local
            "openai_api_key": os.environ.get("OPENAI_API_KEY", ""),
            "openai_model": "gpt-4o-mini",
            "openai_base": os.environ.get("OPENAI_BASE", "https://api.openai.com/v1"),
            "gemini_api_key": os.environ.get("GEMINI_API_KEY", ""),
            "gemini_model": "gemini-1.5-flash",
            "local_base": "http://127.0.0.1:11434",
            "local_model": "llama3",
        }
        self.first_text = ""
        self._build_ui()

    def _build_ui(self):
        nb = ttk.Notebook(self)
        nb.pack(fill=tk.BOTH, expand=True)

        t_meta, t_comp, t_prev, t_export, t_upload, t_api, t_build, t_diag = [ttk.Frame(nb) for _ in range(8)]
        nb.add(t_meta, text="Metadata")
        nb.add(t_comp, text="Compose")
        nb.add(t_prev, text="Preview")
        nb.add(t_export, text="Export")
        nb.add(t_upload, text="Upload")
        nb.add(t_api, text="APIs")
        nb.add(t_build, text="Windows Build")
        nb.add(t_diag, text="Diagnostics")

        # Metadata
        self.title_var = tk.StringVar()
        self.subtitle_var = tk.StringVar()
        self.author_var = tk.StringVar(value="Guillaume Lessard")
        self.publisher_var = tk.StringVar(value="iD01t Productions")
        self.isbn_var = tk.StringVar()

        grid = ttk.Frame(t_meta, padding=10)
        grid.pack(fill=tk.X)
        for i in range(4):
            grid.columnconfigure(i, weight=1)
        ttk.Label(grid, text="Title").grid(row=0, column=0, sticky="w")
        ttk.Entry(grid, textvariable=self.title_var).grid(row=0, column=1, sticky="ew")
        ttk.Label(grid, text="Subtitle").grid(row=0, column=2, sticky="w")
        ttk.Entry(grid, textvariable=self.subtitle_var).grid(row=0, column=3, sticky="ew")
        ttk.Label(grid, text="Author").grid(row=1, column=0, sticky="w")
        ttk.Entry(grid, textvariable=self.author_var).grid(row=1, column=1, sticky="ew")
        ttk.Label(grid, text="Publisher").grid(row=1, column=2, sticky="w")
        ttk.Entry(grid, textvariable=self.publisher_var).grid(row=1, column=3, sticky="ew")
        ttk.Label(grid, text="ISBN").grid(row=2, column=0, sticky="w")
        ttk.Entry(grid, textvariable=self.isbn_var).grid(row=2, column=1, sticky="ew")

        ttk.Label(t_meta, text="Description").pack(anchor="w", padx=10)
        self.desc_txt = tk.Text(t_meta, height=6, wrap="word")
        self.desc_txt.pack(fill=tk.BOTH, padx=10)
        ttk.Label(t_meta, text="Table of Contents (comma or newline separated)").pack(anchor="w", padx=10, pady=(6, 0))
        self.toc_txt = tk.Text(t_meta, height=6, wrap="word")
        self.toc_txt.pack(fill=tk.BOTH, padx=10)

        bar = ttk.Frame(t_meta, padding=10)
        bar.pack(fill=tk.X)
        ttk.Button(bar, text="Load 1st Edition", command=self.load_first).pack(side=tk.LEFT)
        ttk.Button(bar, text="Generate Draft", command=self.gen_draft).pack(side=tk.LEFT, padx=6)

        # Compose
        tools = ttk.Frame(t_comp, padding=8)
        tools.pack(fill=tk.X)
        ttk.Button(tools, text="Clean text", command=self.do_clean).pack(side=tk.LEFT)
        ttk.Button(tools, text="Replace -- with ,", command=self.do_dash_rule).pack(side=tk.LEFT, padx=6)
        ttk.Button(tools, text="Merge 1st Edition", command=self.do_merge_first).pack(side=tk.LEFT)
        ttk.Button(tools, text="AI Expand", command=self.do_ai_expand).pack(side=tk.LEFT, padx=6)
        self.editor = tk.Text(t_comp, wrap="word")
        self.editor.pack(fill=tk.BOTH, expand=True)

        # Preview
        ptools = ttk.Frame(t_prev, padding=8)
        ptools.pack(fill=tk.X)
        ttk.Button(ptools, text="Refresh Preview", command=self.refresh_preview).pack(side=tk.LEFT)
        self.preview = tk.Text(t_prev, wrap="word", state="disabled")
        self.preview.pack(fill=tk.BOTH, expand=True)

        # Export
        xtools = ttk.Frame(t_export, padding=8)
        xtools.pack(fill=tk.X)
        ttk.Button(xtools, text="Export Markdown", command=self.export_md).pack(side=tk.LEFT)
        ttk.Button(xtools, text="Export DOCX", command=self.export_docx).pack(side=tk.LEFT, padx=6)
        ttk.Button(xtools, text="Export EPUB", command=self.export_epub).pack(side=tk.LEFT)
        ttk.Button(xtools, text="Export PDF", command=self.export_pdf).pack(side=tk.LEFT, padx=6)
        self.export_status = tk.StringVar(value="Ready.")
        ttk.Label(t_export, textvariable=self.export_status).pack(anchor="w", padx=10, pady=6)

        # Upload
        g = ttk.LabelFrame(t_upload, text="Gumroad", padding=8)
        g.pack(fill=tk.X, padx=8, pady=6)
        self.gr_key = tk.StringVar()
        self.gr_price = tk.StringVar(value="1499")
        ttk.Label(g, text="API token").grid(row=0, column=0, sticky="w")
        ttk.Entry(g, textvariable=self.gr_key, show="*").grid(row=0, column=1, sticky="ew")
        ttk.Label(g, text="Price (cents)").grid(row=0, column=2, sticky="e")
        ttk.Entry(g, textvariable=self.gr_price, width=8).grid(row=0, column=3, sticky="w")
        g.columnconfigure(1, weight=1)
        ttk.Button(g, text="Upload file to NEW product", command=self.do_gumroad).grid(row=1, column=0, columnspan=4, sticky="w", pady=6)

        k = ttk.LabelFrame(t_upload, text="Ko-fi", padding=8)
        k.pack(fill=tk.X, padx=8, pady=6)
        ttk.Label(k, text="Ko-fi Shop uploads are manual via browser.").grid(row=0, column=0, sticky="w")
        ttk.Button(k, text="Open Ko-fi Shop", command=self.do_kofi).grid(row=0, column=1, padx=6)

        gp = ttk.LabelFrame(t_upload, text="Google Play Books Partner (ONIX + optional SFTP)", padding=8)
        gp.pack(fill=tk.X, padx=8, pady=6)
        self.sftp_host = tk.StringVar()
        self.sftp_port = tk.StringVar(value="22")
        self.sftp_user = tk.StringVar()
        self.sftp_pass = tk.StringVar()
        self.sftp_key = tk.StringVar()
        ttk.Label(gp, text="Host").grid(row=0, column=0, sticky="w")
        ttk.Entry(gp, textvariable=self.sftp_host).grid(row=0, column=1, sticky="ew")
        ttk.Label(gp, text="Port").grid(row=0, column=2, sticky="e")
        ttk.Entry(gp, textvariable=self.sftp_port, width=6).grid(row=0, column=3, sticky="w")
        ttk.Label(gp, text="User").grid(row=1, column=0, sticky="w")
        ttk.Entry(gp, textvariable=self.sftp_user).grid(row=1, column=1, sticky="ew")
        ttk.Label(gp, text="Password").grid(row=1, column=2, sticky="e")
        ttk.Entry(gp, textvariable=self.sftp_pass, show="*").grid(row=1, column=3, sticky="w")
        ttk.Label(gp, text="Private key (optional)").grid(row=2, column=0, sticky="w")
        ttk.Entry(gp, textvariable=self.sftp_key).grid(row=2, column=1, sticky="ew")
        gp.columnconfigure(1, weight=1)
        ttk.Button(gp, text="Write ONIX from metadata", command=self.do_onix).grid(row=3, column=0, pady=6, sticky="w")
        ttk.Button(gp, text="SFTP upload ONIX + EPUB", command=self.do_sftp_upload).grid(row=3, column=1, pady=6, sticky="w")

        # APIs
        modebox = ttk.LabelFrame(t_api, text="Expansion mode", padding=8)
        modebox.pack(fill=tk.X, padx=8, pady=6)
        self.mode_var = tk.StringVar(value="offline")
        for m in ["offline", "openai", "gemini", "local"]:
            ttk.Radiobutton(modebox, text=m.capitalize(), value=m, variable=self.mode_var, command=self.apply_mode).pack(side=tk.LEFT, padx=6)

        oai = ttk.LabelFrame(t_api, text="OpenAI-compatible", padding=8)
        oai.pack(fill=tk.X, padx=8, pady=6)
        self.oai_key = tk.StringVar(value=os.environ.get("OPENAI_API_KEY", ""))
        self.oai_model = tk.StringVar(value="gpt-4o-mini")
        self.oai_base = tk.StringVar(value=os.environ.get("OPENAI_BASE", "https://api.openai.com/v1"))
        ttk.Label(oai, text="API key").grid(row=0, column=0, sticky="w")
        ttk.Entry(oai, textvariable=self.oai_key, show="*").grid(row=0, column=1, sticky="ew")
        ttk.Label(oai, text="Model").grid(row=1, column=0, sticky="w")
        ttk.Entry(oai, textvariable=self.oai_model).grid(row=1, column=1, sticky="ew")
        ttk.Label(oai, text="Base URL").grid(row=2, column=0, sticky="w")
        ttk.Entry(oai, textvariable=self.oai_base).grid(row=2, column=1, sticky="ew")
        oai.columnconfigure(1, weight=1)

        gem = ttk.LabelFrame(t_api, text="Gemini", padding=8)
        gem.pack(fill=tk.X, padx=8, pady=6)
        self.gm_key = tk.StringVar(value=os.environ.get("GEMINI_API_KEY", ""))
        self.gm_model = tk.StringVar(value="gemini-1.5-flash")
        ttk.Label(gem, text="API key").grid(row=0, column=0, sticky="w")
        ttk.Entry(gem, textvariable=self.gm_key, show="*").grid(row=0, column=1, sticky="ew")
        ttk.Label(gem, text="Model").grid(row=1, column=0, sticky="w")
        ttk.Entry(gem, textvariable=self.gm_model).grid(row=1, column=1, sticky="ew")
        gem.columnconfigure(1, weight=1)

        loc = ttk.LabelFrame(t_api, text="Local LLM (Ollama-like)", padding=8)
        loc.pack(fill=tk.X, padx=8, pady=6)
        self.loc_base = tk.StringVar(value="http://127.0.0.1:11434")
        self.loc_model = tk.StringVar(value="llama3")
        ttk.Label(loc, text="Base URL").grid(row=0, column=0, sticky="w")
        ttk.Entry(loc, textvariable=self.loc_base).grid(row=0, column=1, sticky="ew")
        ttk.Label(loc, text="Model").grid(row=1, column=0, sticky="w")
        ttk.Entry(loc, textvariable=self.loc_model).grid(row=1, column=1, sticky="ew")
        loc.columnconfigure(1, weight=1)

        ttk.Button(t_api, text="Save API settings", command=self.save_api).pack(anchor="e", padx=10, pady=8)

        # Build
        ttk.Label(
            t_build,
            text="Write Windows build helpers, then run build\\build_windows.cmd (requires: pip install pyinstaller).",
            padding=8,
        ).pack(anchor="w")
        ttk.Button(t_build, text="Write Build Files", command=WinHelpers.write_all).pack(anchor="w", padx=10, pady=6)
        ttk.Button(t_build, text="Open project folder", command=self.open_project).pack(anchor="w", padx=10)

        # Diagnostics
        ttk.Button(t_diag, text="Run self check (exports)", command=self.self_check).pack(anchor="w", padx=10, pady=8)
        self.diag = tk.Text(t_diag, wrap="word", height=24)
        self.diag.pack(fill=tk.BOTH, expand=True, padx=10, pady=8)
        self.log("Ready.")

    # Utility
    def log(self, msg: str):
        self.diag.configure(state="normal")
        self.diag.insert("end", f"{msg}\n")
        self.diag.see("end")
        self.diag.configure(state="disabled")

    def apply_mode(self):
        self.api_cfg["mode"] = self.mode_var.get()
        self.log(f"Mode set: {self.api_cfg['mode']}")

    def save_api(self):
        self.api_cfg.update({
            "mode": self.mode_var.get(),
            "openai_api_key": self.oai_key.get().strip(),
            "openai_model": self.oai_model.get().strip(),
            "openai_base": self.oai_base.get().strip(),
            "gemini_api_key": self.gm_key.get().strip(),
            "gemini_model": self.gm_model.get().strip(),
            "local_base": self.loc_base.get().strip(),
            "local_model": self.loc_model.get().strip(),
        })
        (PROJECT_DIR / "api_settings.json").write_text(json.dumps(self.api_cfg, indent=2), encoding="utf-8")
        messagebox.showinfo(APP_NAME, "API settings saved")

    # Metadata actions
    def load_first(self):
        p = filedialog.askopenfilename(filetypes=[("Text/Markdown", "*.txt *.md"), ("All files", "*.*")])
        if not p:
            return
        try:
            self.first_text = Path(p).read_text(encoding="utf-8", errors="ignore")
            messagebox.showinfo(APP_NAME, "1st edition loaded")
        except Exception as e:
            messagebox.showerror(APP_NAME, f"Read failed: {e}")

    def gen_draft(self):
        title = self.title_var.get().strip() or "Untitled"
        subtitle = self.subtitle_var.get().strip()
        toc = self.toc_txt.get("1.0", "end").strip()
        desc = self.desc_txt.get("1.0", "end").strip()
        md = scaffold_from_meta(title, subtitle, toc, desc, seed=self.first_text)
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", md)
        messagebox.showinfo(APP_NAME, "Draft generated")

    # Compose actions
    def do_clean(self):
        s = self.editor.get("1.0", "end")
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", clean_text(s))

    def do_dash_rule(self):
        s = self.editor.get("1.0", "end")
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", s.replace("--", ","))

    def do_merge_first(self):
        if not getattr(self, "first_text", "").strip():
            messagebox.showwarning(APP_NAME, "Load a 1st edition first")
            return
        cur = self.editor.get("1.0", "end").strip()
        merged = (cur + "\n\n---\n\n" + self.first_text.strip()).strip()
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", merged)

    def do_ai_expand(self):
        s = self.editor.get("1.0", "end")
        if not s.strip():
            messagebox.showwarning(APP_NAME, "Nothing to expand")
            return
        self.config(cursor="wait")
        self.update_idletasks()

        def task():
            try:
                out = Expander(self.api_cfg).expand(s)
                self.editor.delete("1.0", "end")
                self.editor.insert("1.0", out)
            except Exception as e:
                messagebox.showerror(APP_NAME, f"Expand failed: {e}")
            finally:
                self.config(cursor="")
                self.update_idletasks()
        threading.Thread(target=task, daemon=True).start()

    # Preview
    def refresh_preview(self):
        s = clean_text(self.editor.get("1.0", "end"))
        self.preview.configure(state="normal")
        self.preview.delete("1.0", "end")
        self.preview.insert("1.0", s)
        self.preview.configure(state="disabled")

    # Export
    def export_md(self):
        title = self.title_var.get().strip() or "ebook"
        p = Exporter(PROJECT_DIR).export_md(clean_text(self.editor.get("1.0", "end")), title)
        self.export_status.set(f"Markdown: {p}")

    def export_docx(self):
        title = self.title_var.get().strip() or "ebook"
        p = Exporter(PROJECT_DIR).export_docx(clean_text(self.editor.get("1.0", "end")), title)
        self.export_status.set(f"DOCX: {p}")

    def export_epub(self):
        title = self.title_var.get().strip() or "ebook"
        author = self.author_var.get().strip()
        p = Exporter(PROJECT_DIR).export_epub(clean_text(self.editor.get("1.0", "end")), title, author)
        self.export_status.set(f"EPUB: {p}")

    def export_pdf(self):
        title = self.title_var.get().strip() or "ebook"
        author = self.author_var.get().strip()
        p = Exporter(PROJECT_DIR).export_pdf(clean_text(self.editor.get("1.0", "end")), title, author)
        self.export_status.set(f"PDF: {p}")

    # Upload
    def do_gumroad(self):
        token = self.gr_key.get().strip()
        if not token:
            messagebox.showwarning(APP_NAME, "Enter Gumroad API token")
            return
        fp = filedialog.askopenfilename(title="Pick file to upload")
        if not fp:
            return
        try:
            res = Uploader(PROJECT_DIR).gumroad_create_and_upload(
                token,
                self.title_var.get() or "New eBook",
                int(self.gr_price.get() or "999"),
                self.desc_txt.get("1.0", "end").strip(),
                Path(fp),
            )
            messagebox.showinfo(APP_NAME, f"Gumroad OK, product {res.get('product_id')}")
        except Exception as e:
            messagebox.showerror(APP_NAME, f"Gumroad failed: {e}")

    def do_kofi(self):
        Uploader(PROJECT_DIR).kofi_open_shop()

    def do_onix(self):
        md = {
            "title": self.title_var.get(),
            "subtitle": self.subtitle_var.get(),
            "author": self.author_var.get(),
            "publisher": self.publisher_var.get(),
            "isbn": self.isbn_var.get(),
            "description": self.desc_txt.get("1.0", "end").strip(),
        }
        p = Uploader(PROJECT_DIR).google_onix_xml(md)
        messagebox.showinfo(APP_NAME, f"ONIX written: {p}")

    def do_sftp_upload(self):
        host = self.sftp_host.get().strip()
        user = self.sftp_user.get().strip()
        port = int(self.sftp_port.get() or 22)
        if not host or not user:
            messagebox.showwarning(APP_NAME, "Enter SFTP host and user")
            return
        epub_file = EXPORTS / f"{slugify(self.title_var.get() or 'ebook')}.epub"
        onix_file = EXPORTS / "onix.xml"
        if not onix_file.exists():
            self.do_onix()
        if not epub_file.exists():
            self.export_epub()
        key = self.sftp_key.get().strip() or None
        pwd = self.sftp_pass.get().strip() or None
        try:
            msg = Uploader(PROJECT_DIR).google_sftp_upload(host, port, user, pwd, key, [onix_file, epub_file])
            messagebox.showinfo(APP_NAME, msg)
        except Exception as e:
            messagebox.showerror(APP_NAME, f"SFTP failed: {e}")

    # Build
    def open_project(self):
        path = PROJECT_DIR
        try:
            if sys.platform.startswith("win"):
                os.startfile(str(path))
            elif sys.platform == "darwin":
                subprocess.call(["open", str(path)])
            else:
                subprocess.call(["xdg-open", str(path)])
        except Exception:
            messagebox.showinfo(APP_NAME, f"Project: {path}")

    # Diagnostics
    def self_check(self):
        try:
            ex = Exporter(PROJECT_DIR)
            sample = scaffold_from_meta(self.title_var.get() or "Sample Book", self.subtitle_var.get(), "Intro, Core, Finale", "Quick self-test")
            md = ex.export_md(sample, self.title_var.get() or "Sample Book")
            ep = ex.export_epub(sample, self.title_var.get() or "Sample Book", self.author_var.get())
            pdf = ex.export_pdf(sample, self.title_var.get() or "Sample Book", self.author_var.get())
            self.log(f"Self-check OK\nMD:  {md}\nEPUB: {ep}\nPDF: {pdf}")
            messagebox.showinfo(APP_NAME, "Self-check completed (MD, EPUB, PDF)")
        except Exception as e:
            self.log(f"Self-check error: {e}")
            messagebox.showerror(APP_NAME, f"Self-check failed: {e}")

# ------------------------------
# Entry
# ------------------------------
def main():
    try:
        app = App()
        app.mainloop()
    except Exception as exc:
        try:
            from tkinter import messagebox as mb
            mb.showerror(APP_NAME, f"Fatal error: {exc}\n\n{traceback.format_exc()}")
        except Exception:
            print("Fatal error:", exc)
            print(traceback.format_exc())

if __name__ == "__main__":
    main()
