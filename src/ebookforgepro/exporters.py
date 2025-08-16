import re
import html
import textwrap
from pathlib import Path

# Third-party libraries for exporting
from ebooklib import epub
import markdown2
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas as pdf_canvas

# Local imports
from .core import slugify, APP_ID, EXPORTS

def _esc(s: str) -> str:
    """A tiny HTML escaper."""
    return html.escape(s or "")

class Exporter:
    def __init__(self, proj: Path):
        self.proj = proj
        EXPORTS.mkdir(parents=True, exist_ok=True)

    def export_md(self, text: str, title: str) -> Path:
        p = EXPORTS / f"{slugify(title or 'manuscript')}.md"
        p.write_text(text, encoding="utf-8")
        return p

    def export_docx(self, text: str, title: str) -> Path:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        for line in text.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("> "):
                p = doc.add_paragraph()
                r = p.add_run(line[2:].strip())
                r.italic = True
            else:
                doc.add_paragraph(line)
        p = EXPORTS / f"{slugify(title or 'manuscript')}.docx"
        doc.save(p)
        return p

    def export_epub(self, text: str, title: str, author: str) -> Path:
        book = epub.EpubBook()
        book.set_identifier(APP_ID)
        book.set_title(title or "Untitled")
        book.add_author(author or "")
        book.set_language("en")

        parts = re.split(r'\n## ', text)
        items = []
        for i, part in enumerate(parts):
            if i == 0:
                # Front matter
                content = markdown2.markdown(part)
                c = epub.EpubHtml(title="Introduction", file_name="intro.xhtml", lang="en")
                c.content = f"<h1>{_esc(title)}</h1>{content}"
            else:
                # Chapters
                chapter_title_match = re.match(r'(\d+)\. (.+)', part.splitlines()[0])
                if chapter_title_match:
                    chapter_title = chapter_title_match.group(2)
                    body = "\n".join(part.splitlines()[1:])
                else:
                    chapter_title = f"Chapter {i}"
                    body = part

                c = epub.EpubHtml(title=_esc(chapter_title), file_name=f'chap_{i:02d}.xhtml', lang='en')
                c.content = f"<h2>{_esc(chapter_title)}</h2>" + markdown2.markdown(body)

            book.add_item(c)
            items.append(c)

        book.toc = tuple(items)
        book.spine = ['nav'] + items
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        out_path = EXPORTS / f"{slugify(title or 'ebook')}.epub"
        epub.write_epub(str(out_path), book)
        return out_path

    def export_pdf(self, text: str, title: str, author: str) -> Path:
        out_path = EXPORTS / f"{slugify(title or 'ebook')}.pdf"
        c = pdf_canvas.Canvas(str(out_path), pagesize=LETTER)
        width, height = LETTER
        margin = 72  # 1 inch

        y = height - margin
        c.setFont("Times-Bold", 18)
        c.drawString(margin, y, (title or "Untitled")[:100])
        y -= 24
        if author:
            c.setFont("Times-Roman", 12)
            c.drawString(margin, y, f"By: {author}")
            y -= 36

        c.setFont("Times-Roman", 11)
        plain = re.sub(r'<[^>]+>', '', text) # Basic HTML tag stripping

        for para in plain.split('\n\n'):
            lines = textwrap.wrap(para, width=80)
            for line in lines:
                if y < margin:
                    c.showPage()
                    c.setFont("Times-Roman", 11)
                    y = height - margin
                c.drawString(margin, y, line)
                y -= 14
            y -= 6

        c.save()
        return out_path
